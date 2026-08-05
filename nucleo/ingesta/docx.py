# -*- coding: utf-8 -*-
"""
================================================================================
 INGESTA DE DOCX  -  fragmentacion por estructura logica
 Entregable 4 de 9. Agosto 2026.
================================================================================

Por que esta capa importa mas que el modelo
-------------------------------------------
Un fragmento mal cortado no se arregla con un modelo mejor: si la tabla que
mapea servidores a VLAN y segmentos IP llega partida, ninguna cantidad de
parametros la reconstruye. Aqui se decide el techo de calidad del sistema.

Como se detectan las secciones
------------------------------
No basta con los estilos de titulo de Word. Medido sobre un corpus real de 19
documentos: 5 no usan NINGUN estilo de titulo, y en uno de ellos los titulos de
seccion son tablas de una celda. Por eso hay tres senales, en orden:

  1. estilo Heading de Word
  2. tabla de 1x1 cuyo texto empieza con numeracion  ->  '1 OBJETIVO'
  3. parrafo plano que empieza con numeracion        ->  '1.1 Objetivo'

La 3 es la que rescata secciones que el autor olvido marcar como titulo.

Las tablas de una celda son ambiguas
------------------------------------
En este tipo de documento una tabla 1x1 puede ser tres cosas distintas, y
tratarlas igual arruina el resultado:

    '1 OBJETIVO'              -> titulo de seccion
    'ATENCION: nunca...'      -> callout, pertenece a la seccion actual
    (logo, celda vacia)       -> decorativo, se descarta

Las tablas de datos NUNCA se parten
-----------------------------------
Una tabla es una unidad de sentido. Va entera en su fragmento, con el titulo de
su seccion como contexto, aunque exceda el tamano objetivo. Partirla por llegar
al limite de tokens seria cumplir la metrica destruyendo el dato.

Nada aqui conoce a ningun cliente: los patrones son genericos y los ajustes
especificos entran por la configuracion del tenant.
================================================================================
"""

from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


# =============================================================================
#  PERFIL DE DOCUMENTO  -  como esta ARMADO el documento, no de que trata
# =============================================================================
#  Cada empresa maqueta distinto. Solo en el corpus de referencia hay tres
#  plantillas: una marca las secciones con estilos de Word, otra las mete
#  dentro de tablas de una celda, y una tercera usa estilo de titulo para los
#  items de una lista.
#
#  Por eso NADA de esto puede estar fijo en el nucleo: llega otra empresa con
#  su cuarta plantilla y habria que tocar el motor. El perfil viaja en la
#  configuracion del tenant y aqui solo quedan convenciones genericas del
#  espanol administrativo, como valores por defecto que cualquiera sobrescribe.

# '1 OBJETIVO', '1.1 Objetivo', '2.3.1 Algo'. Exige que despues del numero
# venga una letra: evita confundir '1. Conecta el cable' de una lista numerada
# con un titulo de seccion.
RE_NUMERACION = re.compile(r"^(\d+(?:\.\d+)*)[.\)]?\s+(\S.*)$")

# Referencias cruzadas: 'ver seccion 3.2', 'segun el numeral 4.1'.
RE_REFERENCIA = re.compile(
    r"(?:secci[oó]n|numeral|apartado|punto)\s+(\d+(?:\.\d+)*)", re.I)


@dataclass
class PerfilDocumento:
    """La plantilla documental de UNA empresa. Se define en su config."""

    # Como empieza un recuadro de aviso. Genericos del espanol; una empresa que
    # use otra palabra la agrega y listo.
    marcas_callout: tuple[str, ...] = (
        "⚠", "📝", "❗", "ℹ", "✅", "🔴",
        "atencion", "nota", "importante", "advertencia",
        "precaucion", "recuerde", "ojo")

    # Bloques de pie que se repiten en cada documento. Vectorizarlos mete el
    # mismo ruido en todas las busquedas.
    encabezados_pie: tuple[str, ...] = (
        "registro de cambio", "aprobacion", "elaborado",
        "revisado por", "aprobado por", "control de cambios")

    # Cabeceras de la tabla de metadatos del documento.
    campos_metadatos: tuple[str, ...] = ("codigo", "version", "fecha")

    # Una tabla de una celda con '1 OBJETIVO' abre seccion. Cierto en las
    # plantillas que meten los titulos en tablas; falso en las que no.
    titulo_un_nivel_en_tabla: bool = True

    # Sin estilo de titulo de Word, solo cuenta la numeracion MULTINIVEL
    # ('1.1'). Es lo que separa un titulo de un paso de lista numerada.
    exigir_multinivel_sin_estilo: bool = True

    max_largo_titulo: int = 120

    # Que defectos reportar. 'seccion_vacia' NO viene por defecto: en el corpus
    # de referencia dio 26 falsos positivos y 0 aciertos, porque muchas
    # plantillas usan estilo de titulo para lineas que YA son el contenido.
    # Un detector que grita en falso deja de leerse.
    defectos_a_reportar: tuple[str, ...] = (
        "numeracion_duplicada", "referencia_rota")

    # Deja constancia de las imagenes en el texto del fragmento. No las lee:
    # solo evita que su contenido desaparezca EN SILENCIO.
    anotar_imagenes: bool = True


PERFIL_POR_DEFECTO = PerfilDocumento()


def _sin_acentos(texto: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFD", texto)
                   if unicodedata.category(c) != "Mn")


def _norm(texto: str) -> str:
    return re.sub(r"\s+", " ", _sin_acentos(texto or "")).strip().lower()


# =============================================================================
#  MODELO
# =============================================================================

@dataclass
class Bloque:
    """Un parrafo o una tabla, en el orden en que aparecen en el documento."""
    tipo: str                      # parrafo | tabla
    texto: str = ""
    estilo: str = ""
    filas: list[list[str]] = field(default_factory=list)

    @property
    def dimensiones(self) -> tuple[int, int]:
        return (len(self.filas), len(self.filas[0]) if self.filas else 0)


@dataclass
class Fragmento:
    numero: str                    # '1.4'
    titulo: str
    contenido: str                 # lo que se MUESTRA
    tipo: str                      # texto | tabla
    orden: int
    metadata: dict = field(default_factory=dict)

    def contextualizar(self, doc: "Documento") -> str:
        """
        Antepone la procedencia. Es lo que se VECTORIZA.

        Un fragmento suelto que dice 'se configura en 192.168.1.1' no se
        recupera bien; con 'Del documento G-GO-06 «Guia de configuracion de
        ONT», version 01, seccion 1.4 Configuracion WAN:' delante, si.
        """
        partes = [f"Del documento {doc.codigo}" if doc.codigo else "Del documento"]
        if doc.titulo:
            partes.append(f"«{doc.titulo}»")
        if doc.version:
            partes.append(f"version {doc.version}")
        if self.numero or self.titulo:
            partes.append(f"seccion {self.numero} {self.titulo}".strip())
        return f"{' '.join(partes)}:\n{self.contenido}"


@dataclass
class Documento:
    ruta: Path
    codigo: str = ""
    titulo: str = ""
    version: str = ""
    fecha: str = ""
    # vigente | obsoleto. Un obsoleto se carga pero no se recupera:
    # queda constancia de que existio y de con cual se respondio en su momento.
    estado: str = "vigente"
    fragmentos: list[Fragmento] = field(default_factory=list)
    defectos: list[dict] = field(default_factory=list)


# =============================================================================
#  LECTURA  -  parrafos y tablas EN ORDEN
# =============================================================================

def leer_bloques(ruta: Path,
                 perfil: PerfilDocumento | None = None) -> list[Bloque]:
    """
    python-docx expone document.paragraphs y document.tables por separado, y con
    eso se pierde el orden real. Aqui se recorre el XML del cuerpo para saber
    que tabla va despues de que parrafo — sin eso es imposible saber a que
    seccion pertenece cada tabla.
    """
    perfil = perfil or PERFIL_POR_DEFECTO
    d = docx.Document(str(ruta))
    bloques: list[Bloque] = []
    for hijo in d.element.body.iterchildren():
        if hijo.tag == qn("w:p"):
            p = Paragraph(hijo, d)
            texto = p.text.strip()
            # Una imagen no se lee, pero DEJA CONSTANCIA. Descartarla en
            # silencio hace que el asistente responda con la mitad del
            # procedimiento sin que nadie se entere; con la marca, al menos
            # puede decir "el documento muestra una captura en este punto".
            if perfil.anotar_imagenes and hijo.findall(".//" + qn("w:drawing")):
                texto = (f"{texto}\n[imagen en el documento original]"
                         if texto else "[imagen en el documento original]")
            if texto:
                bloques.append(Bloque("parrafo", texto=texto,
                                      estilo=p.style.name or ""))
        elif hijo.tag == qn("w:tbl"):
            t = Table(hijo, d)
            filas = []
            for r in t.rows:
                fila, anterior = [], None
                for c in r.cells:
                    # Una celda combinada aparece REPETIDA en row.cells: Word la
                    # expone una vez por columna que abarca. No se puede
                    # detectar comparando el texto —dos columnas distintas
                    # pueden tener el mismo valor, como GW y DNS cuando
                    # coinciden— asi que se compara la IDENTIDAD del elemento
                    # XML, que es la unica senal fiable.
                    if anterior is not None and c._tc is anterior:
                        continue
                    fila.append(c.text.strip())
                    anterior = c._tc
                filas.append(fila)
            if any(any(c for c in f) for f in filas):
                bloques.append(Bloque("tabla", filas=filas))
    return bloques


# =============================================================================
#  CLASIFICACION
# =============================================================================

def _es_titulo(bloque: Bloque,
               perfil: PerfilDocumento) -> tuple[str, str] | None:
    """
    Devuelve (numero, titulo) si el bloque abre una seccion.

    El caso dificil es distinguir un titulo de un ITEM DE LISTA. Ambos empiezan
    igual:

        '1.1  Objetivo'                      -> titulo de seccion
        '1.  Conecta tu celular al puerto'   -> paso de un procedimiento

    Confundirlos no es cosmetico: los pasos quedan repartidos en secciones
    falsas, se inventan numeraciones duplicadas que no existen, y el texto de
    esos pasos deja de analizarse (con lo que se pierden las referencias
    cruzadas que contengan).

    La senal que los separa es la PROFUNDIDAD de la numeracion. Un titulo sin
    estilo de Word tiene que ser multinivel ('1.1', '2.3.1'); las listas usan un
    solo nivel. Los titulos de un solo nivel existen, pero en este tipo de
    documento vienen dentro de una tabla, que ya es senal suficiente por si
    misma.
    """
    p = perfil
    if bloque.tipo == "parrafo":
        texto = bloque.texto
        estilo = bloque.estilo or ""
        m = RE_NUMERACION.match(texto)

        if "Heading" in estilo or "Titulo" in estilo or "Título" in estilo:
            return (m.group(1), m.group(2)) if m else ("", texto)

        # Un parrafo con formato de lista nunca es un titulo, aunque este
        # numerado.
        if "List" in estilo or "Lista" in estilo:
            return None

        if not m or len(texto) >= p.max_largo_titulo:
            return None
        if p.exigir_multinivel_sin_estilo and "." not in m.group(1):
            return None
        return m.group(1), m.group(2)

    filas, cols = bloque.dimensiones
    if filas != 1:
        return None

    # Tabla 1x1: '1 OBJETIVO'. Aqui si puede valer un solo nivel.
    if cols == 1:
        if not p.titulo_un_nivel_en_tabla:
            return None
        texto = bloque.filas[0][0].strip()
        if _es_callout(texto, p):
            return None
        m = RE_NUMERACION.match(texto)
        if m and len(texto) < p.max_largo_titulo:
            return m.group(1), m.group(2)
        return None

    # Tabla 1x2 con el numero en una celda y el titulo en la otra:
    #   | 1 | CONFIGURACION DE ONT |
    if cols == 2:
        izq, der = (c.strip() for c in bloque.filas[0][:2])
        if re.fullmatch(r"\d+(?:\.\d+)*", izq) and 0 < len(der) < 150:
            return izq, der
    return None


def _es_callout(texto: str, perfil: PerfilDocumento) -> bool:
    t = _norm(texto)
    return any(t.startswith(_norm(m)) for m in perfil.marcas_callout)


def _es_pie(bloque: Bloque, perfil: PerfilDocumento) -> bool:
    if bloque.tipo != "tabla" or not bloque.filas:
        return False
    primera = _norm(" ".join(bloque.filas[0]))
    return any(_norm(e) in primera for e in perfil.encabezados_pie)


def _es_metadatos(bloque: Bloque, perfil: PerfilDocumento) -> bool:
    if bloque.tipo != "tabla" or len(bloque.filas) < 2:
        return False
    cabecera = [_norm(c) for c in bloque.filas[0]]
    return sum(1 for c in perfil.campos_metadatos if _norm(c) in cabecera) >= 2


def _tabla_a_texto(filas: list[list[str]]) -> str:
    """
    Markdown, no texto plano: conserva la relacion fila-columna. Que el modelo
    reciba 'Servidor 1 | VLAN 300 | 172.16.4.0/22' y no seis palabras sueltas
    es la diferencia entre responder bien y confundir dos servidores.
    """
    if not filas:
        return ""
    # Las celdas combinadas ya vienen deshechas desde leer_bloques().
    limpias = filas
    ancho = max(len(f) for f in limpias)
    if ancho == 0:
        return ""

    def _linea(f: list[str]) -> str:
        return "| " + " | ".join(f + [""] * (ancho - len(f))) + " |"

    lineas = [_linea(limpias[0]), "|" + "|".join([" --- "] * ancho) + "|"]
    for f in limpias[1:]:
        # Fila que quedo en una sola celda dentro de una tabla ancha: era una
        # nota que abarcaba todo el ancho. Se saca de la tabla para no fingir
        # una columna que no existe.
        if len(f) == 1 and ancho > 1:
            lineas.append("")
            lineas.append(f[0])
        else:
            lineas.append(_linea(f))
    return "\n".join(lineas)


def _extraer_metadatos(bloque: Bloque) -> dict:
    cab = [_norm(c) for c in bloque.filas[0]]
    val = bloque.filas[1]
    out = {}
    for campo in ("codigo", "version", "fecha"):
        if campo in cab:
            i = cab.index(campo)
            if i < len(val) and val[i].strip():
                out[campo] = val[i].strip()
    return out


# =============================================================================
#  FRAGMENTACION
# =============================================================================

def procesar(ruta: Path,
             perfil: PerfilDocumento | None = None,
             max_tokens: int = 500) -> Documento:
    """
    Fragmenta un DOCX segun el PERFIL de la empresa que lo produjo.

    El perfil describe la plantilla —como marca las secciones, que recuadros
    son avisos, que bloques son pie de pagina—. Sin el, cada empresa nueva
    obligaria a tocar este archivo, y los criterios de una se aplicarian a los
    documentos de la otra.
    """
    perfil = perfil or PERFIL_POR_DEFECTO
    bloques = leer_bloques(ruta, perfil)
    doc = Documento(ruta=ruta)

    # --- metadatos y titulo, de las primeras tablas ---
    for b in bloques[:6]:
        if _es_metadatos(b, perfil):
            m = _extraer_metadatos(b)
            doc.codigo = m.get("codigo", "")
            doc.version = m.get("version", "")
            doc.fecha = m.get("fecha", "")
        elif b.tipo == "tabla" and b.dimensiones == (1, 1) and not doc.titulo:
            # Solo la primera linea: la celda del titulo suele traer debajo el
            # area y la razon social, y eso repetido en cada fragmento
            # contextualizado es ruido en todas las busquedas.
            t = b.filas[0][0].strip().splitlines()[0].strip() if b.filas[0][0].strip() else ""
            if t and not RE_NUMERACION.match(t) and len(t) < 150:
                doc.titulo = t
    if not doc.codigo:
        # Ultimo recurso: el nombre del archivo suele empezar por el codigo.
        m = re.match(r"^([A-Z]-[A-Z]{2}-\d+)", ruta.stem)
        doc.codigo = m.group(1) if m else ruta.stem[:40]
    if not doc.titulo:
        doc.titulo = ruta.stem.replace("_", " ")

    # --- recorrido: se acumula texto por seccion, las tablas van aparte ---
    num_actual, tit_actual = "", "(preambulo)"
    buffer: list[str] = []
    orden = 0
    secciones_vistas: dict[str, int] = {}
    referencias: list[str] = []

    def cerrar_texto():
        nonlocal buffer, orden
        texto = "\n".join(buffer).strip()
        buffer = []
        if not texto:
            return
        for trozo in _partir_texto(texto, max_tokens):
            orden += 1
            doc.fragmentos.append(Fragmento(
                numero=num_actual, titulo=tit_actual, contenido=trozo,
                tipo="texto", orden=orden,
                metadata={"seccion": num_actual}))

    for b in bloques:
        if _es_pie(b, perfil):
            continue                                   # boilerplate

        titulo = _es_titulo(b, perfil)
        if titulo:
            cerrar_texto()
            num_actual, tit_actual = titulo
            if num_actual:
                secciones_vistas[num_actual] = secciones_vistas.get(num_actual, 0) + 1
            continue

        if b.tipo == "parrafo":
            buffer.append(b.texto)
            referencias += RE_REFERENCIA.findall(b.texto)
            continue

        # --- tablas ---
        if _es_metadatos(b, perfil):
            continue
        filas, cols = b.dimensiones
        if (filas, cols) == (1, 1):
            texto = b.filas[0][0].strip()
            if _es_callout(texto, perfil):
                buffer.append(texto)                   # pertenece a la seccion
                referencias += RE_REFERENCIA.findall(texto)
            continue                                   # decorativa: se descarta

        # Tabla de datos: fragmento propio, ENTERA, sin importar su tamano.
        cerrar_texto()
        orden += 1
        doc.fragmentos.append(Fragmento(
            numero=num_actual, titulo=tit_actual,
            contenido=_tabla_a_texto(b.filas), tipo="tabla", orden=orden,
            metadata={"seccion": num_actual, "filas": filas, "columnas": cols}))

    cerrar_texto()

    # --- defectos ---
    activos = set(perfil.defectos_a_reportar)

    if "numeracion_duplicada" in activos:
        for numero, veces in sorted(secciones_vistas.items()):
            if veces > 1:
                doc.defectos.append({
                    "tipo": "numeracion_duplicada", "seccion": numero,
                    "detalle": f"la seccion {numero} aparece {veces} veces"})

    if "referencia_rota" in activos:
        for ref in sorted(set(referencias)):
            if ref not in secciones_vistas:
                doc.defectos.append({
                    "tipo": "referencia_rota", "seccion": ref,
                    "detalle": f"se remite a la seccion {ref}, que no existe"})

    if "seccion_vacia" not in activos:
        return doc

    # "Seccion vacia" solo es un defecto util si de verdad falta contenido. Dos
    # salvedades, ambas medidas contra un corpus real:
    #
    #   - una seccion CONTENEDORA no esta vacia: su contenido son sus hijas
    #   - un titulo que ES una frase completa ya es el contenido. Se ve en los
    #     perfiles de cargo, donde el autor uso estilo de titulo para los items
    #     de una lista: "4.3 Procurar el cuidado integral de su salud."
    #
    # Sin estas dos, el detector reporto 21 defectos falsos contra 7 reales, y
    # un detector asi no se vuelve a mirar.
    con_contenido = {f.numero for f in doc.fragmentos}
    titulos = {}
    for b in bloques:
        t = _es_titulo(b, perfil)
        if t and t[0]:
            titulos[t[0]] = t[1]

    for n in sorted(secciones_vistas):
        if n in con_contenido:
            continue
        if any(otra.startswith(f"{n}.") for otra in secciones_vistas):
            continue                                   # tiene subsecciones
        titulo = titulos.get(n, "")
        if len(titulo) > 60 or titulo.rstrip().endswith("."):
            continue                                   # el titulo es el contenido
        doc.defectos.append({"tipo": "seccion_vacia", "seccion": n,
                             "detalle": f"la seccion {n} no tiene contenido"})
    return doc


def _partir_texto(texto: str, max_tokens: int) -> list[str]:
    """
    Parte por parrafos, nunca a media frase. ~4 caracteres por token en espanol.
    Un parrafo mas largo que el limite se deja entero: cortarlo perderia mas de
    lo que ahorra.
    """
    limite = max_tokens * 4
    if len(texto) <= limite:
        return [texto]
    trozos, actual = [], ""
    for parrafo in texto.split("\n"):
        if actual and len(actual) + len(parrafo) + 1 > limite:
            trozos.append(actual)
            actual = parrafo
        else:
            actual = f"{actual}\n{parrafo}" if actual else parrafo
    if actual:
        trozos.append(actual)
    return trozos
